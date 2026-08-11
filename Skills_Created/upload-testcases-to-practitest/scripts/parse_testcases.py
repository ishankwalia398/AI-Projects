#!/usr/bin/env python3
"""
Parse test case files (CSV, Excel, PDF, Word, Markdown) into a structured format
suitable for uploading to PractiTest.

This script handles common test case file formats and returns a list of test case
dictionaries with fields: name, description, steps, custom_fields, etc.

Usage:
    python parse_testcases.py <file_path> [--sheet <sheet_name>]

Output: JSON array of test cases to stdout
"""

import argparse
import json
import sys
from pathlib import Path
from typing import List, Dict, Any, Optional

def parse_csv(file_path: str) -> List[Dict[str, Any]]:
    """Parse CSV file into test cases."""
    import pandas as pd
    df = pd.read_csv(file_path)
    return _parse_dataframe(df)

def parse_excel(file_path: str, sheet_name: Optional[str] = None) -> List[Dict[str, Any]]:
    """Parse Excel file into test cases."""
    import pandas as pd

    # If no sheet specified, list available sheets
    if sheet_name is None:
        xls = pd.ExcelFile(file_path)
        if len(xls.sheet_names) > 1:
            print(f"Multiple sheets found: {xls.sheet_names}", file=sys.stderr)
            print(f"Using first sheet: {xls.sheet_names[0]}", file=sys.stderr)
        sheet_name = xls.sheet_names[0]

    df = pd.read_excel(file_path, sheet_name=sheet_name)
    return _parse_dataframe(df)

def _parse_dataframe(df: 'pd.DataFrame') -> List[Dict[str, Any]]:
    """
    Parse a DataFrame into test cases.

    Handles two formats:
    1. One row per test case (each row is a complete test case)
    2. One row per step (test case info on first row, steps on subsequent rows)
    """
    import pandas as pd  # Re-import to ensure pandas is available in function scope

    # Clean column names (strip whitespace, convert to lowercase for matching)
    df.columns = [str(col).strip() for col in df.columns]

    # Detect format by looking for step-related columns
    step_indicators = ['step#', 'step_#', 'step', 'stepno', 'step_no', 'step_number']
    has_step_column = any(col.lower().replace(' ', '_') in step_indicators for col in df.columns)

    if has_step_column:
        return _parse_multi_row_format(df)
    else:
        return _parse_single_row_format(df)

def _parse_single_row_format(df: 'pd.DataFrame') -> List[Dict[str, Any]]:
    """Parse format where each row is a complete test case."""
    import pandas as pd
    test_cases = []

    for idx, row in df.iterrows():
        tc = {
            'name': _extract_field(row, ['tc_name', 'test_name', 'testcase', 'name', 'title']),
            'description': _extract_field(row, ['tc_description', 'description', 'desc', 'summary']),
            'preconditions': _extract_field(row, ['preconditions', 'precondition', 'prerequisites']),
            'steps': [],
            'custom_fields': {}
        }

        # Try to extract steps from a single field (if present)
        steps_field = _extract_field(row, ['steps', 'test_steps', 'actions'])
        if steps_field:
            tc['steps'] = _parse_steps_from_text(steps_field)

        # Extract expected result (might be a single field for the whole test case)
        expected_result = _extract_field(row, ['expected_result', 'expected', 'result'])
        if expected_result and not tc['steps']:
            tc['steps'].append({
                'name': 'Execute',
                'description': steps_field or tc['description'],
                'expected_results': expected_result
            })

        # Priority
        priority = _extract_field(row, ['priority', 'pri'])
        if priority:
            tc['priority'] = priority

        # Status
        status = _extract_field(row, ['status', 'state'])
        if status:
            tc['status'] = status

        # Add any other columns as custom fields
        for col in df.columns:
            col_lower = col.lower().replace(' ', '_')
            if col_lower not in ['tc_name', 'test_name', 'testcase', 'name', 'title',
                                  'tc_description', 'description', 'desc', 'summary',
                                  'preconditions', 'precondition', 'prerequisites',
                                  'steps', 'test_steps', 'actions', 'expected_result',
                                  'expected', 'result', 'priority', 'pri', 'status', 'state']:
                value = row[col]
                if pd.notna(value):
                    tc['custom_fields'][col] = str(value)

        # Only add if we have at least a name
        if tc['name']:
            test_cases.append(tc)
        else:
            print(f"Warning: Row {idx+1} missing test case name, skipping", file=sys.stderr)

    return test_cases

def _parse_multi_row_format(df: 'pd.DataFrame') -> List[Dict[str, Any]]:
    """
    Parse format where each row is a step, and test case info is on the first row of each group.

    Expected structure:
    - TC_ID or TC_Name column identifies which test case the row belongs to
    - First row of each group has test case-level info (name, description)
    - All rows have step-level info (step name, step description, expected result)
    """
    import pandas as pd

    # Find the grouping column (TC_ID or similar)
    group_col = None
    for col in df.columns:
        col_lower = col.lower().replace(' ', '_')
        if col_lower in ['tc_id', 'testcase_id', 'test_id', 'id', 'tc_name', 'test_name']:
            group_col = col
            break

    if not group_col:
        print("Warning: Could not find TC_ID or TC_Name column for grouping. Treating as single-row format.", file=sys.stderr)
        return _parse_single_row_format(df)

    test_cases = []
    grouped = df.groupby(group_col, sort=False)

    for tc_id, group in grouped:
        if pd.isna(tc_id):
            continue

        # Test case-level info from first row
        first_row = group.iloc[0]
        tc = {
            'name': _extract_field(first_row, ['tc_name', 'test_name', 'testcase', 'scenario_name', 'name']),
            'description': _extract_field(first_row, ['tc_description', 'description', 'desc']),
            'preconditions': None,
            'steps': [],
            'custom_fields': {}
        }

        # If no explicit name field, use the group key
        if not tc['name']:
            tc['name'] = str(tc_id)

        # Extract steps from all rows
        for idx, row in group.iterrows():
            step = {
                'name': _extract_field(row, ['step_name', 'action', 'step']),
                'description': _extract_field(row, ['step_description', 'step_desc', 'step_details']),
                'expected_results': _extract_field(row, ['expected_result', 'expected', 'expected_results'])
            }

            # If step name is "Precondition", treat it specially
            if step['name'] and step['name'].lower() in ['precondition', 'preconditions', 'prerequisite']:
                tc['preconditions'] = step['description'] or step['expected_results']
            else:
                # Only add steps with at least a name or description
                if step['name'] or step['description']:
                    tc['steps'].append(step)

        test_cases.append(tc)

    return test_cases

def _extract_field(row: 'pd.Series', possible_names: List[str]) -> Optional[str]:
    """
    Extract a field value from a row, trying multiple possible column names.
    Returns None if not found or if value is NaN.
    """
    import pandas as pd

    for name in possible_names:
        for col in row.index:
            if col.lower().replace(' ', '_').replace('-', '_') == name.lower():
                value = row[col]
                if pd.notna(value):
                    return str(value).strip()
    return None

def _parse_steps_from_text(text: str) -> List[Dict[str, str]]:
    """
    Parse steps from a text block (e.g., numbered list).
    Returns a list of step dictionaries.
    """
    # Simple heuristic: split on newlines and look for numbered patterns
    lines = text.split('\n')
    steps = []

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Check if line starts with a number (e.g., "1.", "1)", "1:")
        import re
        match = re.match(r'^(\d+)[.):\s]+(.+)', line)
        if match:
            step_text = match.group(2).strip()
            steps.append({
                'name': f'Step {match.group(1)}',
                'description': step_text,
                'expected_results': None
            })
        else:
            # Not numbered, treat as a single step description
            steps.append({
                'name': 'Step',
                'description': line,
                'expected_results': None
            })

    return steps if steps else None

def parse_pdf(file_path: str) -> List[Dict[str, Any]]:
    """
    Parse PDF file into test cases.

    PDFs are tricky because they don't have a clear structure. This function
    tries to extract text and parse it, but results may vary depending on the PDF format.
    """
    try:
        import PyPDF2
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            text = ''
            for page in reader.pages:
                text += page.extract_text()

        # Try to parse as a simple list of test cases (very basic)
        # In practice, PDFs with tables should be converted to CSV/Excel first
        print("Warning: PDF parsing is limited. Consider converting to CSV/Excel for better results.", file=sys.stderr)
        return _parse_text_format(text)
    except ImportError:
        print("Error: PyPDF2 not installed. Run: pip install PyPDF2", file=sys.stderr)
        sys.exit(1)

def parse_word(file_path: str) -> List[Dict[str, Any]]:
    """
    Parse Word (.docx) file into test cases.

    Looks for tables and paragraphs that might contain test case information.
    """
    try:
        from docx import Document
        doc = Document(file_path)

        test_cases = []

        # Try to extract from tables first
        for table in doc.tables:
            # Assume first row is headers
            if len(table.rows) < 2:
                continue

            headers = [cell.text.strip().lower() for cell in table.rows[0].cells]

            for row in table.rows[1:]:
                row_data = {headers[i]: cell.text.strip() for i, cell in enumerate(row.cells)}

                tc = {
                    'name': _extract_field_from_dict(row_data, ['name', 'test_name', 'tc_name', 'title']),
                    'description': _extract_field_from_dict(row_data, ['description', 'desc', 'summary']),
                    'steps': [],
                    'custom_fields': {}
                }

                # Extract steps if present
                steps_text = _extract_field_from_dict(row_data, ['steps', 'test_steps', 'actions'])
                if steps_text:
                    tc['steps'] = _parse_steps_from_text(steps_text)

                if tc['name']:
                    test_cases.append(tc)

        # If no tables found, try parsing paragraphs (less structured)
        if not test_cases:
            text = '\n'.join([para.text for para in doc.paragraphs])
            test_cases = _parse_text_format(text)

        return test_cases
    except ImportError:
        print("Error: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
        sys.exit(1)

def _extract_field_from_dict(data: Dict[str, str], possible_keys: List[str]) -> Optional[str]:
    """Extract a field from a dictionary, trying multiple possible keys."""
    for key in possible_keys:
        for actual_key in data.keys():
            if actual_key.lower().replace(' ', '_').replace('-', '_') == key.lower():
                value = data[actual_key]
                if value and value.strip():
                    return value.strip()
    return None

def parse_markdown(file_path: str) -> List[Dict[str, Any]]:
    """
    Parse Markdown file into test cases.

    Looks for tables or structured lists.
    """
    with open(file_path, 'r', encoding='utf-8') as f:
        text = f.read()

    return _parse_text_format(text)

def _parse_text_format(text: str) -> List[Dict[str, Any]]:
    """
    Parse unstructured text into test cases.

    This is a fallback for PDFs, plain text, or Markdown without clear structure.
    Results may be incomplete.
    """
    # Very basic heuristic: look for lines that might be test case names
    # (e.g., "Test Case: ...", "TC_001: ...", etc.)
    import re

    test_cases = []
    current_tc = None

    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Check if this line starts a new test case
        match = re.match(r'^(test case|tc|testcase)[:\s#]+(.+)', line, re.IGNORECASE)
        if match:
            if current_tc:
                test_cases.append(current_tc)
            current_tc = {
                'name': match.group(2).strip(),
                'description': '',
                'steps': [],
                'custom_fields': {}
            }
        elif current_tc:
            # Accumulate description
            current_tc['description'] += line + ' '

    if current_tc:
        test_cases.append(current_tc)

    return test_cases

def main():
    parser = argparse.ArgumentParser(description='Parse test case files into JSON format')
    parser.add_argument('file_path', help='Path to the test case file')
    parser.add_argument('--sheet', help='Sheet name for Excel files (optional)')
    parser.add_argument('--output', help='Output file path (default: stdout)')

    args = parser.parse_args()

    file_path = Path(args.file_path)
    if not file_path.exists():
        print(f"Error: File not found: {file_path}", file=sys.stderr)
        sys.exit(1)

    # Detect file type
    suffix = file_path.suffix.lower()

    if suffix == '.csv':
        test_cases = parse_csv(str(file_path))
    elif suffix in ['.xlsx', '.xls']:
        test_cases = parse_excel(str(file_path), args.sheet)
    elif suffix == '.pdf':
        test_cases = parse_pdf(str(file_path))
    elif suffix in ['.docx', '.doc']:
        test_cases = parse_word(str(file_path))
    elif suffix in ['.md', '.markdown']:
        test_cases = parse_markdown(str(file_path))
    else:
        print(f"Error: Unsupported file type: {suffix}", file=sys.stderr)
        sys.exit(1)

    # Output as JSON
    output = json.dumps(test_cases, indent=2, ensure_ascii=False)

    if args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            f.write(output)
    else:
        print(output)

    # Summary to stderr
    print(f"\nParsed {len(test_cases)} test cases from {file_path.name}", file=sys.stderr)
    if test_cases:
        print(f"Sample: {test_cases[0]['name']}", file=sys.stderr)
        total_steps = sum(len(tc['steps']) for tc in test_cases)
        print(f"Total steps: {total_steps}", file=sys.stderr)

if __name__ == '__main__':
    main()
