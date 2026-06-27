"""
Generate Word (.docx) and PDF files from structured test plan content.

Usage:
    python generate_documents.py --content-file content.json --output-dir ./test-plans --name "project-name"

The content.json should have this structure:
{
    "title": "Test Plan Title",
    "sections": [
        {
            "heading": "Introduction",
            "level": 1,
            "content": "Paragraph text here",
            "type": "paragraph"
        },
        {
            "heading": "Test Cases",
            "level": 1,
            "type": "table",
            "headers": ["Test Case", "Description", "Priority"],
            "rows": [
                ["TC-01: Login", "Verify user can log in", "High"],
                ["TC-02: Logout", "Verify user can log out", "Medium"]
            ]
        }
    ]
}
"""

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx library not installed. Install with: pip install python-docx")
    sys.exit(1)


def create_docx(content_data, output_path):
    """Create a Word document from structured content."""
    doc = Document()

    # Set document title
    title = doc.add_heading(content_data.get("title", "Test Plan"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Process each section
    for section in content_data.get("sections", []):
        section_type = section.get("type", "paragraph")

        if section_type == "heading":
            level = section.get("level", 1)
            heading_text = section.get("heading", "")
            doc.add_heading(heading_text, level=level)

        elif section_type == "paragraph":
            heading = section.get("heading")
            if heading:
                doc.add_heading(heading, level=section.get("level", 2))

            content = section.get("content", "")
            if content:
                p = doc.add_paragraph(content)

        elif section_type == "table":
            heading = section.get("heading")
            if heading:
                doc.add_heading(heading, level=section.get("level", 2))

            headers = section.get("headers", [])
            rows = section.get("rows", [])

            if headers and rows:
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                table.style = 'Light Grid Accent 1'

                # Add headers
                header_cells = table.rows[0].cells
                for idx, header in enumerate(headers):
                    header_cells[idx].text = str(header)
                    # Make header bold
                    for paragraph in header_cells[idx].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

                # Add data rows
                for row_idx, row_data in enumerate(rows, start=1):
                    row_cells = table.rows[row_idx].cells
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < len(row_cells):
                            row_cells[col_idx].text = str(cell_data)

        elif section_type == "list":
            heading = section.get("heading")
            if heading:
                doc.add_heading(heading, level=section.get("level", 2))

            items = section.get("items", [])
            for item in items:
                doc.add_paragraph(str(item), style='List Bullet')

    # Save the document
    doc.save(output_path)
    print(f"✓ Created Word document: {output_path}")
    return output_path


def convert_to_pdf(docx_path, pdf_path):
    """Convert the Word document to PDF."""
    try:
        # Try docx2pdf first (works on Windows and Mac)
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        print(f"✓ Created PDF document: {pdf_path}")
        return True
    except ImportError:
        # Try pypandoc as fallback
        try:
            import pypandoc
            pypandoc.convert_file(docx_path, 'pdf', outputfile=str(pdf_path))
            print(f"✓ Created PDF document: {pdf_path}")
            return True
        except ImportError:
            print("Warning: Neither docx2pdf nor pypandoc is installed. PDF conversion skipped.")
            print("Install with: pip install docx2pdf (Windows/Mac) or pip install pypandoc")
            return False
        except Exception as e:
            print(f"Warning: PDF conversion failed: {e}")
            return False


def main():
    parser = argparse.ArgumentParser(description="Generate Word and PDF documents from structured content")
    parser.add_argument("--content-file", required=True, help="Path to JSON file with content structure")
    parser.add_argument("--output-dir", required=True, help="Directory to save output files")
    parser.add_argument("--name", required=True, help="Base name for output files (e.g., 'project-name')")

    args = parser.parse_args()

    # Read content file
    content_path = Path(args.content_file)
    if not content_path.exists():
        print(f"Error: Content file not found: {args.content_file}")
        sys.exit(1)

    with open(content_path, 'r', encoding='utf-8') as f:
        content_data = json.load(f)

    # Create output directory
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    # Generate files
    docx_path = output_dir / f"{args.name}-test-plan.docx"
    pdf_path = output_dir / f"{args.name}-test-plan.pdf"

    # Create Word document
    create_docx(content_data, docx_path)

    # Convert to PDF
    convert_to_pdf(docx_path, pdf_path)

    print(f"\n✓ Document generation complete!")
    print(f"  - Word: {docx_path}")
    print(f"  - PDF: {pdf_path}")


if __name__ == "__main__":
    main()
